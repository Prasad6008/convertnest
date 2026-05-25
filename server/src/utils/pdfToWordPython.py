import sys
import os
import tempfile
import fitz
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_SECTION


def editable_pdf_to_docx(input_pdf, output_docx):
    cv = Converter(input_pdf)

    try:
        cv.convert(output_docx, start=0, end=None)
    finally:
        cv.close()


def set_section_size_and_margins(section, width_in, height_in):
    section.page_width = Inches(width_in)
    section.page_height = Inches(height_in)

    # Keep tiny margins to avoid Word pushing image to next blank page
    section.top_margin = Inches(0.05)
    section.bottom_margin = Inches(0.05)
    section.left_margin = Inches(0.05)
    section.right_margin = Inches(0.05)


def exact_visual_pdf_to_docx(input_pdf, output_docx):
    pdf = fitz.open(input_pdf)
    doc = Document()

    if len(pdf) == 0:
        doc.add_paragraph("No pages found in PDF.")
        doc.save(output_docx)
        return

    # Remove paragraph spacing globally
    normal_style = doc.styles["Normal"]
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.line_spacing = 1

    with tempfile.TemporaryDirectory() as temp_dir:
        for index, page in enumerate(pdf):
            page_width_in = page.rect.width / 72
            page_height_in = page.rect.height / 72

            if index == 0:
                section = doc.sections[0]
            else:
                section = doc.add_section(WD_SECTION.NEW_PAGE)

            set_section_size_and_margins(section, page_width_in, page_height_in)

            available_width = page_width_in - 0.10
            available_height = page_height_in - 0.10

            # Render page as high-quality image
            zoom = 2.0
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)

            image_path = os.path.join(temp_dir, f"page-{index + 1}.png")
            pix.save(image_path)

            image_ratio = pix.width / pix.height
            page_ratio = available_width / available_height

            # Fit inside page without overflowing
            if image_ratio > page_ratio:
                final_width = available_width
                final_height = final_width / image_ratio
            else:
                final_height = available_height
                final_width = final_height * image_ratio

            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1

            run = paragraph.add_run()
            run.add_picture(
                image_path,
                width=Inches(final_width),
                height=Inches(final_height)
            )

    pdf.close()
    doc.save(output_docx)


if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python pdfToWordPython.py input.pdf output.docx mode", file=sys.stderr)
        sys.exit(1)

    input_pdf = sys.argv[1]
    output_docx = sys.argv[2]
    mode = sys.argv[3]

    try:
        if mode == "exact":
            exact_visual_pdf_to_docx(input_pdf, output_docx)
        else:
            editable_pdf_to_docx(input_pdf, output_docx)

        print("PDF to Word conversion completed")
    except Exception as error:
        print(str(error), file=sys.stderr)
        sys.exit(1)